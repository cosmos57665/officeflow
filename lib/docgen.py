"""Document generation helpers: Word minutes and PDF certificates."""
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def minutes_docx(data: dict, out_path):
    """Write a professional meeting-minutes Word document.

    Expected keys in data (all optional): title, date, attendees (list),
    summary (str or list of paragraphs), action_items (list of dicts with
    task / owner / deadline).
    """
    doc = Document()

    heading = doc.add_heading(str(data.get("title", "Meeting Minutes")), level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Date: {data.get('date', date.today().isoformat())}").italic = True

    doc.add_heading("Attendees", level=1)
    attendees = data.get("attendees") or []
    doc.add_paragraph(", ".join(str(a) for a in attendees) if attendees else "Not recorded")

    doc.add_heading("Summary", level=1)
    summary = data.get("summary") or "No summary available."
    for para in summary if isinstance(summary, list) else [summary]:
        doc.add_paragraph(str(para))

    doc.add_heading("Action Items", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for cell, label in zip(table.rows[0].cells, ("Task", "Owner", "Deadline")):
        cell.paragraphs[0].add_run(label).bold = True
    for item in data.get("action_items") or []:
        row = table.add_row().cells
        row[0].text = str(item.get("task", ""))
        row[1].text = str(item.get("owner", ""))
        row[2].text = str(item.get("deadline", ""))

    doc.save(str(out_path))


def _latin1(text: str) -> str:
    # fpdf2 core fonts only cover latin-1; replace anything else so we never crash
    return str(text).encode("latin-1", "replace").decode("latin-1")


def certificate_pdf(name: str, remark: str, out_path):
    """Write a clean A4 certificate with border, heading, name, remark and signature line."""
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()

    # double border
    pdf.set_line_width(1.0)
    pdf.rect(10, 10, 190, 277)
    pdf.set_line_width(0.3)
    pdf.rect(13, 13, 184, 271)

    pdf.set_y(55)
    pdf.set_font("Helvetica", "B", 28)
    pdf.cell(0, 15, "Certificate of Achievement", align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(14)
    pdf.set_font("Helvetica", "", 13)
    pdf.cell(0, 8, "This certificate is proudly presented to", align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(8)
    pdf.set_font("Helvetica", "B", 34)
    pdf.cell(0, 18, _latin1(name), align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(12)
    pdf.set_font("Helvetica", "", 12)
    pdf.set_x(35)
    pdf.multi_cell(140, 7, _latin1(remark), align="C")

    # footer: date on the left, signature line on the right
    y = 245
    pdf.set_font("Helvetica", "", 12)
    pdf.text(30, y, f"Date: {date.today().strftime('%d %B %Y')}")
    pdf.line(125, y, 180, y)
    pdf.set_font("Helvetica", "I", 10)
    pdf.text(142, y + 6, "Signature")

    pdf.output(str(out_path))
